import taos
from docx import Document
from taos import connect
import pandas as pd
import numpy as np
from docx.shared import Inches
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from openpyxl import Workbook
from openpyxl import load_workbook
import os
from clean import GPSDataFilter2
import time
'''
v1.2
1、新增结构振动、地震动、船舶撞击分析
2、桥梁编码与桥梁名称、传感器编号一一对应
3、新增数据清洗功能
4、振动没加.any because mei清洗
5、vie没加try
6、高频分页读取数据，处理函数那里略有不同

'''


def read_data(statement):
    conn: taos.TaosConnection = taos.connect(host="jkjc1",
                                             user="root",
                                             password="taosdata",
                                             database="db_jkjc",
                                             port=6030)

    result = conn.query(statement)
    result_list = result.fetch_all()  # 将查询结果转换为列表
    conn.close()
    return result_list

    # # 连接到TDengine数据库
    # conn = connect(host='jkjc1', user='root', password='taosdata', database='db_jkjc', port=6030)
    # chunksize = 200000
    # offset = 0
    # # 初始化一个空的列表来存储所有块
    # chunks = []
    # # 创建游标对象
    # cursor = conn.cursor()
    #
    # while True:
    #     sql = statement.format(chunksize, offset)
    #     cursor.execute(sql)
    #
    #     # 获取结果并转换为 DataFrame
    #     rows = cursor.fetchall()
    #     if not rows:
    #         break
    #
    #     df_chunk = pd.DataFrame(rows, columns=['ts', 'val'])
    #     chunks.append(df_chunk)
    #     offset += chunksize
    #
    # # 将所有数据块合并成一个DataFrame
    # df = pd.concat(chunks, ignore_index=True)
    #
    # # 关闭游标和连接
    # cursor.close()
    # conn.close()
    #
    # return df


def data_calculate(data):
    maximum = np.nanmax(data)
    minimum = np.nanmin(data)
    mean = np.nanmean(data)
    chazhi = maximum-minimum
    result = [round(maximum, 4), round(minimum, 4), round(chazhi, 4)]
    return result


def analysis_nd(bridge_num, sensor_num, sensor_position, time_1, time_2, bridge_name, season_num, type, unit_1):
    query_words = ('select ts,val from ' + '`' + bridge_num + '-' + sensor_num + '`' +
                   ' where ts >= ' + "'" + time_1 + "'" + ' and ts <= ' + "'" + time_2 + "'" + 'limit {} offset {}')

    data = read_data(query_words)
    time = data['ts'].tolist()
    data_nd = data['val'].tolist()
    data_nd = GPSDataFilter2(data_nd, 50, 0.5, 360, 1.5)
    path = fig_plot(time, data_nd, sensor_num, sensor_position, bridge_name, season_num, type, unit_1)
    if data_nd.any():
        tongjizhi = data_calculate(data_nd)
    else:
        tongjizhi = ['nan', 'nan', 'nan']

    return tongjizhi, path


def analysis_crk(bridge_num, sensor_num, sensor_position, time_1, time_2, bridge_name, season_num, type, unit_1):
    query_words = ('select ts,val from ' + '`' + bridge_num + '-' + sensor_num + '`' +
                   ' where ts >= ' + "'" + time_1 + "'" + ' and ts <= ' + "'" + time_2 + "'")
    try:
        data = read_data(query_words)
        time = [t[0] for t in data]
        data_crk_origin = [t[1] for t in data]
        data_crk = GPSDataFilter2(data_crk_origin, 0.5, 0.05, 3600, 1.5)
        path = fig_plot(time, data_crk, sensor_num, sensor_position, bridge_name, season_num, type, unit_1)
        if data_crk.any():
            tongjizhi = data_calculate(data_crk)
        else:
            tongjizhi = ['nan', 'nan', 'nan']
    except Exception as e:
        print(f"Error occurred: {e}")
        tongjizhi = ['999', '999', '999']
        path_1 = r'E:\project\auto_report'  # 前缀地址，允许自定义
        path = os.path.join(path_1, bridge_name, type, season_num)
    return tongjizhi, path


def analysis_dis01(bridge_num, sensor_num, sensor_position, time_1, time_2, bridge_name, season_num, type, unit_1):
    query_words = ('select ts,val from ' + '`' + bridge_num + '-' + sensor_num + '`' +
                   ' where ts >= ' + "'" + time_1 + "'" + ' and ts <= ' + "'" + time_2 + "'")
    try:
        data = read_data(query_words)
        time = [t[0] for t in data]
        data_dis01_origin = [t[1] for t in data]
        data_dis01 = GPSDataFilter2(data_dis01_origin, 100, 1, 3600, 1.5)
        path = fig_plot(time, data_dis01, sensor_num, sensor_position, bridge_name, season_num, type, unit_1)
        if data_dis01.any():
            tongjizhi = data_calculate(data_dis01)
        else:
            tongjizhi = ['nan', 'nan', 'nan']
    except Exception as e:
        print(f"Error occurred: {e}")
        tongjizhi = ['999', '999', '999']
        path_1 = r'E:\project\auto_report'  # 前缀地址，允许自定义
        path = os.path.join(path_1, bridge_name, type, season_num)
    return tongjizhi, path


def analysis_dis02(bridge_num, sensor_num, sensor_position, time_1, time_2, bridge_name, season_num, type, unit_1):
    query_words = ('select ts,val from ' + '`' + bridge_num + '-' + sensor_num + '`' +
                   ' where ts >= ' + "'" + time_1 + "'" + ' and ts <= ' + "'" + time_2 + "'")
    try:
        data = read_data(query_words)
        time = [t[0] for t in data]
        data_dis02_origin = [t[1] for t in data]
        data_dis02 = GPSDataFilter2(data_dis02_origin, 100, 1, 3600, 1.5)
        path = fig_plot(time, data_dis02, sensor_num, sensor_position, bridge_name, season_num, type, unit_1)
        if data_dis02.any():
            tongjizhi = data_calculate(data_dis02)
        else:
            tongjizhi = ['nan', 'nan', 'nan']
    except Exception as e:
        print(f"Error occurred: {e}")
        tongjizhi = ['999', '999', '999']
        path_1 = r'E:\project\auto_report'  # 前缀地址，允许自定义
        path = os.path.join(path_1, bridge_name, type, season_num)
    return tongjizhi, path


def analysis_vib(bridge_num, sensor_num, sensor_position, time_1, time_2, bridge_name, season_num, type, unit_1):
    query_words = ('select ts,val from ' + '`' + bridge_num + '-' + sensor_num + '`' +
                   ' where ts >= ' + "'" + time_1 + "'" + ' and ts <= ' + "'" + time_2 + "'" + 'limit {} offset {}')
    data = read_data(query_words)
    time = data['ts'].tolist()
    data_vib = data['val'].tolist()
    path = fig_plot(time, data_vib, sensor_num, sensor_position, bridge_name, season_num, type, unit_1)
    if data_vib:
        tongjizhi = data_calculate(data_vib)
    else:
        tongjizhi = ['nan', 'nan', 'nan']
    return tongjizhi, path


def analysis_vie(bridge_num, sensor_num, sensor_position, time_1, time_2, bridge_name, season_num, type, unit_1):
    query_words = ('select ts,val from ' + '`' + bridge_num + '-' + sensor_num + '`' +
                   ' where ts >= ' + "'" + time_1 + "'" + ' and ts <= ' + "'" + time_2 + "'" + 'limit {} offset {}')
    try:
        data = read_data(query_words)
        time = data['ts'].tolist()
        data_vie = data['val'].tolist()
        path = fig_plot(time, data_vie, sensor_num, sensor_position, bridge_name, season_num, type, unit_1)
        if data_vie:
            tongjizhi = data_calculate(data_vie)
        else:
            tongjizhi = ['nan', 'nan', 'nan']
    except Exception as e:
        print(f"Error occurred: {e}")
        tongjizhi = ['999', '999', '999']
        path_1 = r'E:\project\auto_report'  # 前缀地址，允许自定义
        path = os.path.join(path_1, bridge_name, type, season_num)
    return tongjizhi, path


def analysis_tmp(bridge_num, sensor_num, sensor_position, time_1, time_2, bridge_name, season_num, type, unit_1):
    query_words = ('select ts,val from ' + '`' + bridge_num + '-' + sensor_num + '`' +
                   ' where ts >= ' + "'" + time_1 + "'" + ' and ts <= ' + "'" + time_2 + "'")
    try:
        data = read_data(query_words)
        time = [t[0] for t in data]
        data_tmp = [t[1] for t in data]
        data_tmp = GPSDataFilter2(data_tmp, 100, 10, 3600, 1.5)
        path = fig_plot(time, data_tmp, sensor_num, sensor_position, bridge_name, season_num, type, unit_1)
        if data_tmp.any():
            tongjizhi = data_calculate(data_tmp)
        else:
            tongjizhi = ['nan', 'nan', 'nan']
    except Exception as e:
        print(f"Error occurred: {e}")
        tongjizhi = ['999', '999', '999']
        path_1 = r'E:\project\auto_report'  # 前缀地址，允许自定义
        path = os.path.join(path_1, bridge_name, type, season_num)
    return tongjizhi, path


def analysis_rhs(bridge_num, sensor_num, sensor_position, time_1, time_2, bridge_name, season_num, type_1, type_2, unit_1, unit_2):
    query_words = ('select ts,val1,val2 from ' + '`' + bridge_num + '-' + sensor_num + '`' +
                   ' where ts >= ' + "'" + time_1 + "'" + ' and ts <= ' + "'" + time_2 + "'")
    try:
        data = read_data(query_words)
        time = [t[0] for t in data]
        data_rts = [t[1] for t in data]
        data_rhs = [t[2] for t in data]
        data_rts = GPSDataFilter2(data_rts, 100, 3, 10000, 1.5)
        data_rhs = GPSDataFilter2(data_rhs, 100, 5, 20000, 1.5)
        path_1 = fig_plot(time, data_rts, sensor_num, sensor_position, bridge_name, season_num, type_1, unit_1)
        path_2 = fig_plot(time, data_rhs, sensor_num, sensor_position, bridge_name, season_num, type_2, unit_2)
        if data_rts.any():
            tongjizhi_1 = data_calculate(data_rts)
        else:
            tongjizhi_1 = ['nan', 'nan', 'nan']
            print(f'{sensor_num}: RTS data is void')
        if data_rhs.any():
            tongjizhi_2 = data_calculate(data_rhs)
        else:
            tongjizhi_2 = ['nan', 'nan', 'nan']
            print(f'{sensor_num}: RHS data is void')
    except Exception as e:
        print(f"Error occurred: {e}")
        tongjizhi_1 = ['999', '999', '999']
        tongjizhi_2 = ['999', '999', '999']
        path_0 = r'E:\project\auto_report'  # 前缀地址，允许自定义
        path_1 = os.path.join(path_0, bridge_name, type_1, season_num)
        path_2 = os.path.join(path_0, bridge_name, type_2, season_num)
    return tongjizhi_1, tongjizhi_2, path_1, path_2


def analysis_rsg(bridge_num, sensor_num, sensor_position, time_1, time_2, bridge_name, season_num, type_1, type_2, unit_1, unit_2):
    query_words = ('select ts,val1,val2,val3 from ' + '`' + bridge_num + '-' + sensor_num + '`' +
                   ' where ts >= ' + "'" + time_1 + "'" + ' and ts <= ' + "'" + time_2 + "'")
    try:
        data = read_data(query_words)
        time = [t[0] for t in data]
        data_ybw = [t[2] for t in data]
        data_yb = [t[3] for t in data]
        data_ybw = GPSDataFilter2(data_ybw, 100, 1, 36000, 1.5)
        data_yb = GPSDataFilter2(data_yb, 1000, 5, 36000, 1.5)
        path_1 = fig_plot(time, data_ybw, sensor_num, sensor_position, bridge_name, season_num, type_1, unit_1)
        path_2 = fig_plot(time, data_yb, sensor_num, sensor_position, bridge_name, season_num, type_2, unit_2)
        if data_ybw.any():
            tongjizhi_1 = data_calculate(data_ybw)
        else:
            tongjizhi_1 = ['nan', 'nan', 'nan']
            print(f'{sensor_num}: YB data is void')
        if data_yb.any():
            tongjizhi_2 = data_calculate(data_yb)
        else:
            tongjizhi_2 = ['nan', 'nan', 'nan']
            print(f'{sensor_num}: YBw data is void')
    except Exception as e:
        print(f"Error occurred: {e}")
        tongjizhi_1 = ['999', '999', '999']
        tongjizhi_2 = ['999', '999', '999']
        path_0 = r'E:\project\auto_report'  # 前缀地址，允许自定义
        path_1 = os.path.join(path_0, bridge_name, type_1, season_num)
        path_2 = os.path.join(path_0, bridge_name, type_2, season_num)
    return tongjizhi_1, tongjizhi_2, path_1, path_2


def find_paragraph_by_keyword(doc, keyword):
    for paragraph in doc.paragraphs:
        # 查找包含指定关键词的标题段落
        if paragraph.style.name == 'Heading 2' and keyword in paragraph.text:
            return paragraph


def add_data_to_table_after_keyword(doc, keyword, data, pos):
    # 查找包含指定关键词的标题段落
    target_paragraph = find_paragraph_by_keyword(doc, keyword)
    if target_paragraph:
        found_target = False
        for element in doc.element.body:
            if found_target:
                if element.tag.endswith('tbl'):
                    table = element
                    break
            if element.tag.endswith('p') and element.text == target_paragraph.text:
                found_target = True
        else:
            print('Target table not found')
            raise ValueError("Table not found after the specified keyword")

        # 找到表格并将数据添加到表格的后三列
        table = Table(table, doc)
        for row_index, row_data in enumerate(data, start=1):  # 从表格的第二行开始替换数据
            row_cells = table.rows[row_index].cells  # 获取各行的所有单元格
            for i, cell_data in enumerate(row_data):  # 遍历行中的每个单元格数据
                row_cells[-(pos + len(row_data) - i)].text = str(cell_data)  # 填充单元格数据
                row_cells[-(pos + len(row_data) - i)].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 设置表格单元格中第一个段落的对齐方式为居中


def fig_plot(ts, val, sensor, sensor_position, bridge_name, season_num, type, unit_1):
    b_with = 1  # 边框宽度
    # 字体宏观设置
    plt.rcParams['font.sans-serif'] = ['Microsoft YaHei']  # 使用微软雅黑字体
    plt.rcParams['axes.unicode_minus'] = False  # 正常显示负号
    # 设置图片大小
    plt.figure(figsize=(10, 4), dpi=100)
    ax = plt.gca()  # 获取边框
    # 设置边框
    ax.spines['bottom'].set_linewidth(b_with)  # 图框下边
    ax.spines['left'].set_linewidth(b_with)  # 图框左边
    ax.spines['top'].set_linewidth(b_with)  # 图框上边
    ax.spines['right'].set_linewidth(b_with)  # 图框右边
    # 画图，添加图例
    plt.plot(ts, val)
    # plt.legend(loc="best")
    # 坐标轴标题设置
    plt.xlabel("时间", fontdict={'family': 'Microsoft YaHei', 'size': 10}, labelpad=5)
    plt.ylabel(type + unit_1, fontdict={'family': 'Microsoft YaHei', 'size': 10}, labelpad=5)
    # 刻度标签参数
    plt.xticks(fontproperties='Microsoft YaHei', size=10)
    plt.yticks(fontproperties='Microsoft YaHei', size=10)
    # 刻度线的大小长短粗细
    plt.tick_params(axis="both", which="major", direction="in", width=1, length=5, pad=5)
    # 刻度标签自定义
    # plt.yticks(np.arange(min(val), max(val) + 1, (max(val) - min(val)) / 10))

    # 设置日期格式
    date_format = mdates.DateFormatter('%m/%d')  # %Y-%m-%d %H:%M:%S  %m/%d
    ax.xaxis.set_major_formatter(date_format)

    # 设置横坐标为等间距
    ax.xaxis.set_major_locator(mdates.AutoDateLocator())

    # 标题及网格线
    plt.title(sensor + type + '时程数据', fontproperties='Microsoft YaHei', size=12)
    plt.grid(True, linestyle='--', alpha=0.5)
    # plt.show()

    # 创建子文件夹
    path_1 = r'E:\project\auto_report'  # 前缀地址，允许自定义
    # 创建目标文件夹路径
    path_destination = os.path.join(path_1, bridge_name, type, season_num)
    os.makedirs(path_destination, exist_ok=True)
    # 生成并保存图片
    file_path = os.path.join(path_destination, f'{sensor}.png')
    plt.savefig(file_path)
    plt.close()

    return path_destination


def add_picture_to_doc_after_table(doc, keyword, path):
    # 查找包含指定关键词的标题段落
    target_paragraph = find_paragraph_by_keyword(doc, keyword)
    if target_paragraph:
        found_target = False
        for element in doc.element.body:
            if found_target:
                if element.tag.endswith('tbl'):
                    table = element
                    break
            if element.tag.endswith('p') and element.text == target_paragraph.text:
                found_target = True
        else:
            raise ValueError("Table not found after the specified keyword")

        # 在表格后添加图片
        table_idx = doc.element.body.index(table)  # 获取表格在文档主体中的索引
        supported_formats = ['.jpg', '.jpeg', '.png', '.gif', '.bmp']

        image_files = [f for f in os.listdir(path) if os.path.splitext(f)[1].lower() in supported_formats]

        for i, image_file in enumerate(image_files):
            if os.path.splitext(image_file)[1].lower() not in supported_formats:
                print(f"Unsupported image format: {image_file}")
                continue

            new_paragraph = doc.add_paragraph()  # 创建一个新段落
            run = new_paragraph.add_run()  # 创建一个新的文本运行对象

            try:
                run.add_picture(os.path.join(path, image_file), width=Inches(6.0))
            except Exception as e:
                print(f"Error adding picture {image_file}: {e}")
                continue

            # 将新段落插入到表格后
            doc.element.body.insert(table_idx + i * 2 + 1, new_paragraph._element)  # 将新段落插入到表格后
            # 插入编号标题
            caption_paragraph = doc.add_paragraph()
            caption_paragraph.add_run(f"({i + 1}) {image_file}")
            # 将编号标题段落插入到图片后
            doc.element.body.insert(table_idx + i * 2 + 2, caption_paragraph._element)


def save_data_to_excel(data, path):
    if not os.path.exists(path):
        os.makedirs(path)
    else:
        # 创建一个新的Workbook对象
        wb = Workbook()
        ws = wb.active  # 获取默认的sheet

        # 定义标题行和标题列
        title_row = ["SENSOR", "POSITION", "MAX", "MIN", "DIFF"]
        title_column = ["Row1", "Row2", "Row3"]

        # 写入标题行（在第1行）
        for col_num, title in enumerate(title_row, start=2):  # 从第2列开始写入
            ws.cell(row=1, column=col_num, value=title)

        # 定义起始行和列（数据从第2行第4列开始）
        start_row = 2
        start_column = 4

        # 将数据写入指定的行和列
        for i, row in enumerate(data):
            for j, value in enumerate(row):
                ws.cell(row=start_row + i, column=start_column + j, value=value)
        # 保存到Excel文件
        file_path = os.path.join(path, 'output.xlsx')
        wb.save(file_path)


def save_sensor_info_to_excel(data, path):
    # 指定excel文件路径
    file_path = os.path.join(path, 'output.xlsx')

    # 打开Excel工作簿
    wb = load_workbook(file_path)

    # 选择一个工作表（例如：第一个工作表）
    ws = wb.active

    # 定义起始行和列（数据从第2行第2列开始）
    start_row = 2
    start_column = 2

    # 将数据写入指定的行和列
    for i, row in enumerate(data):
        for j, value in enumerate(row):
            ws.cell(row=start_row + i, column=start_column + j, value=value)
    # 保存到Excel文件
    wb.save(file_path)


if __name__ == "__main__":
    bridge = ['G204望虞河大桥']
    '''
    'G204望虞河大桥', 'G204张家港河大桥', 'S228望虞河大桥', 'S228张家港大桥', 'S359望虞河大桥', 'S359杨园互通（含匝道F/H)', 'S359杨园互通F匝道',
              'S359杨园互通H匝道', 'S359尤泾河桥', 'S359张家港大桥'
    '''
    '''
    'G104归径河桥', 'G104钟张运河大桥', 'G312北兴塘大桥', 'G346石牌港桥', 'S230城东港桥', 'S230沙墩港桥', 'S240孟津河桥',
              'S259黄桥', 'S263钮家大桥', 'S263徐舍东大桥', 'S340东青河大桥', 'S342南山桥', 'S342融通港桥', 'S342塘南河桥',
              'S342西氿大桥', 'S342走马塘河桥'
    '''
    season = 'Q2'
    time_start = '2024-04-01 00:00:00'
    time_end = '2024-06-30 23:59:59'
    # sensor_type = 'aaa'
    # sensor = 'hhh'
    df = pd.read_excel(r'E:\project\auto_report\code\sensorinfo_sz.xlsx', sheet_name='BRIDGE_TEST_SELFCHECK.T_BRIDGE')
    # 过滤数据，选取所需的列
    filtered_data = df[df['BRIDGENAME'].isin(bridge)][['FOREIGN_KEY', 'SENSOR_SUB_TYPE_NAME', 'SENSOR_CODE', 'POSITION', 'BRIDGENAME']]
    bridge_number = filtered_data['FOREIGN_KEY'].to_list()

    for bridge_name in bridge:
        bridge_data = filtered_data[filtered_data['BRIDGENAME'] == bridge_name]
        grouped = bridge_data.groupby('SENSOR_SUB_TYPE_NAME')
        for sensor_type, group in grouped:
            data_list = []
            data_list_1 = []
            sensor_list = []
            position_list = []
            sensor_data_list = []
            keyword = '666'
            keyword_1 = '666'
            path = '666'
            path_1 = '666'
            path_2 = '666'
            for sensor_code, position, bridge_number in zip(group['SENSOR_CODE'], group['POSITION'], group['FOREIGN_KEY']):
                start_time = time.time()  # 循环开始前获取当前时间戳
                if sensor_type == '结构裂缝':
                    unit = '(mm)'
                    data_single, path = analysis_crk(bridge_number, sensor_code, position, time_start, time_end, bridge_name, season, sensor_type, unit)
                    data_list.append(data_single)
                    sensor_list.append(sensor_code)
                    position_list.append(position)
                    print(f'{sensor_code}: done')
                    keyword = "混凝土结构裂缝监测"  # 指定关键词
                elif sensor_type == '结构振动':
                    base_unit = "m/s"
                    superscript_two = "\u00B2"
                    unit = f"{base_unit}{superscript_two}"  # 格式化字符串
                    data_single, path = analysis_vib(bridge_number, sensor_code, position, time_start, time_end,
                                                     bridge_name, season, sensor_type, unit)
                    data_list.append(data_single)
                    sensor_list.append(sensor_code)
                    position_list.append(position)
                    print(f'{sensor_code}: done')
                    keyword = "主梁竖向振动监测"  # 指定关键词
                elif sensor_type == '船舶撞击/地震':
                    base_unit = "m/s"
                    superscript_two = "\u00B2"
                    unit = f"{base_unit}{superscript_two}"  # 格式化字符串
                    data_single, path = analysis_vib(bridge_number, sensor_code, position, time_start, time_end,
                                                     bridge_name, season, sensor_type, unit)
                    data_list.append(data_single)
                    sensor_list.append(sensor_code)
                    position_list.append(position)
                    print(f'{sensor_code}: done')
                    keyword = "地震、船撞监测"  # 指定关键词
                elif sensor_type == '桥面温度':
                    unit = '(℃)'
                    data_single, path = analysis_tmp(bridge_number, sensor_code, position, time_start, time_end, bridge_name, season, sensor_type, unit)
                    data_list.append(data_single)
                    sensor_list.append(sensor_code)
                    position_list.append(position)
                    print(f'{sensor_code}: done')
                    keyword = "桥面铺装层温度监测"  # 指定关键词
                elif sensor_type == '梁端纵向位移':
                    unit = '(mm)'
                    data_single, path = analysis_dis01(bridge_number, sensor_code, position, time_start, time_end, bridge_name, season, sensor_type, unit)
                    data_list.append(data_single)
                    sensor_list.append(sensor_code)
                    position_list.append(position)
                    print(f'{sensor_code}: done')
                    keyword = "梁端纵向位移监测"  # 指定关键词
                elif sensor_type == '支座位移':
                    unit = '(mm)'
                    data_single, path = analysis_dis02(bridge_number, sensor_code, position, time_start, time_end, bridge_name, season, sensor_type, unit)
                    data_list.append(data_single)
                    sensor_list.append(sensor_code)
                    position_list.append(position)
                    print(f'{sensor_code}: done')
                    keyword = "支座位移监测"  # 指定关键词
                elif sensor_type == '桥面温度':
                    unit = '(℃)'
                    data_single, path = analysis_tmp(bridge_number, sensor_code, position, time_start, time_end, bridge_name, season, sensor_type, unit)
                    data_list.append(data_single)
                    sensor_list.append(sensor_code)
                    position_list.append(position)
                    print(f'{sensor_code}: done')
                    keyword = "桥面铺装层温度监测"  # 指定关键词
                elif sensor_type == '主梁竖向位移':
                    unit = '(mm)'
                    data_single, path = analysis_nd(bridge_number, sensor_code, position, time_start, time_end, bridge_name, season, sensor_type, unit)
                    data_list.append(data_single)
                    sensor_list.append(sensor_code)
                    position_list.append(position)
                    print(f'{sensor_code}: done')
                    keyword = "主梁竖向位移监测"  # 指定关键词
                elif sensor_type == '环境温湿度':
                    type_1 = '环境温度'
                    type_2 = '环境湿度'
                    unit_1 = '(℃)'
                    unit_2 = '(%)'
                    data_single_1, data_single_2, path_1, path_2 = analysis_rhs(bridge_number, sensor_code, position, time_start, time_end, bridge_name, season, type_1, type_2, unit_1, unit_2)
                    data_list.append(data_single_1)
                    data_list_1.append(data_single_2)
                    sensor_list.append(sensor_code)
                    position_list.append(position)
                    print(f'{sensor_code}: done')
                    keyword = "环境温度、湿度监测"  # 指定关键词
                elif sensor_type == '应变/温度':
                    type_1 = '应变温度'
                    type_2 = '应变'
                    unit_1 = '(℃)'
                    unit_2 = '(uε)'
                    data_single_1, data_single_2, path_1, path_2 = analysis_rsg(bridge_number, sensor_code, position, time_start, time_end, bridge_name, season, type_1, type_2, unit_1, unit_2)
                    data_list.append(data_single_1)
                    data_list_1.append(data_single_2)
                    sensor_list.append(sensor_code)
                    position_list.append(position)
                    print(f'{sensor_code}: done')
                    keyword = "结构温度监测"  # 指定关键词
                    keyword_1 = "关键截面应变监测"  # 指定关键词
                else:
                    print(f'{sensor_code}: Not defined')
                end_time = time.time()  # 循环结束后再次获取当前时间戳
                duration = end_time - start_time  # 计算耗时
                print(f"Loop took {duration:.2f} seconds.")  # 打印耗时

            try:
                if sensor_type == '应变/温度':
                    sensor_info_list = list(zip(sensor_list, position_list))
                    save_data_to_excel(data_list, path_1)
                    save_data_to_excel(data_list_1, path_2)
                    save_sensor_info_to_excel(sensor_info_list, path_1)
                    save_sensor_info_to_excel(sensor_info_list, path_2)
                    # print(sensor_data_list)
                    # doc = Document(r'E:\project\auto_report\code\01.塘南河桥24年4月数据分析报告.docx')  # 读取 Word 文档
                    # add_data_to_table_after_keyword(doc, keyword, data_list, 0)
                    # add_data_to_table_after_keyword(doc, keyword_1, data_list_1, 0)
                    # add_data_to_table_after_keyword(doc, keyword, sensor_info_list, 3)
                    # add_data_to_table_after_keyword(doc, keyword_1, sensor_info_list, 3)
                    # add_picture_to_doc_after_table(doc, keyword, path_1)
                    # add_picture_to_doc_after_table(doc, keyword_1, path_2)
                    # doc.save(r'E:\project\auto_report\code\01.塘南河桥24年4月数据分析报告.docx')  # 保存修改后的文档
                    print(f'{sensor_type}: Solution is done')
                elif sensor_type == '环境温湿度':
                    sensor_info_list = list(zip(sensor_list, position_list))
                    save_data_to_excel(data_list, path_1)
                    save_data_to_excel(data_list_1, path_2)
                    save_sensor_info_to_excel(sensor_info_list, path_1)
                    save_sensor_info_to_excel(sensor_info_list, path_2)
                    data_list.extend(data_list_1)  # 环境温湿度合并列表，其他的也适用，但应变/温度不适用
                    sensor_info_list.extend(sensor_info_list)
                    # doc = Document(r'E:\project\auto_report\code\01.塘南河桥24年4月数据分析报告.docx')  # 读取 Word 文档
                    # add_data_to_table_after_keyword(doc, keyword, data_list, 0)
                    # add_data_to_table_after_keyword(doc, keyword, sensor_info_list, 3)
                    # add_picture_to_doc_after_table(doc, keyword, path_1)
                    # add_picture_to_doc_after_table(doc, keyword, path_2)
                    # doc.save(r'E:\project\auto_report\code\01.塘南河桥24年4月数据分析报告.docx')  # 保存修改后的文档
                    print(f'{sensor_type}: Solution is done')
                else:
                    sensor_info_list = list(zip(sensor_list, position_list))
                    save_data_to_excel(data_list, path)
                    save_sensor_info_to_excel(sensor_info_list, path)
                    data_list.extend(data_list_1)  # 环境温湿度合并列表，其他的也适用，但应变/温度不适用
                    # doc = Document(r'E:\project\auto_report\code\01.塘南河桥24年4月数据分析报告.docx')  # 读取 Word 文档
                    # add_data_to_table_after_keyword(doc, keyword, data_list, 0)
                    # add_data_to_table_after_keyword(doc, keyword, sensor_info_list, 3)
                    # add_picture_to_doc_after_table(doc, keyword, path)
                    # doc.save(r'E:\project\auto_report\code\01.塘南河桥24年4月数据分析报告.docx')  # 保存修改后的文档
                    print(f'{sensor_type}: Solution is done')
            except Exception as e:
                print(f"{e}")

