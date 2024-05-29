import taos
import numpy as np
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
'''
读取单个传感器时间内的数据，绘制图表并添加到word指定位置
'''


def read_data(statement):
    conn: taos.TaosConnection = taos.connect(host="jkjc1",
                                             user="root",
                                             password="taosdata",
                                             database="db_jkjc",
                                             port=6030)

    result = conn.query(statement)
    result_list = result.fetch_all()  # 将查询结果转换为列表
    conn.close()
    return result_list


def data_calculate(data):
    maximum = np.nanmax(data)
    minimum = np.nanmin(data)
    mean = np.nanmean(data)
    chazhi = maximum-minimum
    result = [round(maximum, 4), round(minimum, 4), round(chazhi, 4)]
    return result


def analysis_nd(bridge_num, sensor_num, time_1, time_2, bridge_name, season_num):
    query_words = ('select ts,val from ' + '`' + bridge_num + '-' + sensor_num + '`' +
                   ' where ts >= ' + "'" + time_1 + "'" + ' and ts <= ' + "'" + time_2 + "'")
    data = read_data(query_words)
    time = [t[0] for t in data]
    data_nd = [t[1] for t in data]
    tongjizhi = data_calculate(data_nd)
    fig_plot(time, data_nd, sensor_num, bridge_name, season_num)
    return tongjizhi


def fig_plot(ts, val, sensor, bridge_name, season_num):
    b_with = 1  # 边框宽度
    # 字体宏观设置
    plt.rcParams['font.sans-serif'] = ['Microsoft YaHei']  # 使用微软雅黑字体
    plt.rcParams['axes.unicode_minus'] = False  # 正常显示负号
    # 设置图片大小
    plt.figure(figsize=(10, 4), dpi=100)
    ax = plt.gca()  # 获取边框
    # 设置边框
    ax.spines['bottom'].set_linewidth(b_with)  # 图框下边
    ax.spines['left'].set_linewidth(b_with)  # 图框左边
    ax.spines['top'].set_linewidth(b_with)  # 图框上边
    ax.spines['right'].set_linewidth(b_with)  # 图框右边
    # 画图，添加图例
    plt.plot(ts, val)
    # plt.legend(loc="best")
    # 坐标轴标题设置
    plt.xlabel("时间", fontdict={'family': 'Microsoft YaHei', 'size': 10}, labelpad=5)
    plt.ylabel("主梁竖向位移(mm)", fontdict={'family': 'Microsoft YaHei', 'size': 10}, labelpad=5)
    # 刻度标签参数
    plt.xticks(fontproperties='Microsoft YaHei', size=10)
    plt.yticks(fontproperties='Microsoft YaHei', size=10)
    # 刻度线的大小长短粗细
    plt.tick_params(axis="both", which="major", direction="in", width=1, length=5, pad=5)
    # 刻度标签自定义
    plt.yticks(np.arange(min(val), max(val) + 1, (max(val) - min(val)) / 10))

    # 设置日期格式
    date_format = mdates.DateFormatter('%H:%M:%S')
    ax.xaxis.set_major_formatter(date_format)

    # 标题及网格线
    plt.title('最重车过桥期间' + sensor + '主梁竖向位移波动趋势', fontproperties='Microsoft YaHei', size=12)
    plt.grid(True, linestyle='--', alpha=0.5)
    # plt.show()
    plt.savefig(f'{bridge_name}/{season_num}/{sensor}.png')
    plt.close()


def find_paragraph_by_keyword(doc, keyword):
    for paragraph in doc.paragraphs:
        # 查找包含指定关键词的标题段落
        if paragraph.style.name == 'Heading 2' and keyword in paragraph.text:
            return paragraph


def add_picture_to_doc_after_table(doc, keyword):
    # 查找包含指定关键词的标题段落
    target_paragraph = find_paragraph_by_keyword(doc, keyword)
    if target_paragraph:
        found_target = False
        for element in doc.element.body:
            if found_target:
                if element.tag.endswith('tbl'):
                    table = element
                    break
            if element.tag.endswith('p') and element.text == target_paragraph.text:
                found_target = True
        else:
            raise ValueError("Table not found after the specified keyword")

        # 在表格后添加图片
        table_idx = doc.element.body.index(table)  # 获取表格在文档主体中的索引
        for i in range(1):
            new_paragraph = doc.add_paragraph()  # 创建一个新段落
            run = new_paragraph.add_run()  # 创建一个新的文本运行对象
            run.add_picture(f'figure_{i}.png', width=Inches(6.0))  # 在运行对象中插入图片
            # 将新段落插入到表格后
            doc.element.body.insert(table_idx + 1, new_paragraph._element)  # 将新段落插入到表格后

            # 保存修改后的文档
            doc.save('modified_document.docx')


if __name__ == "__main__":
    bridge = ['S263徐舍东大桥']
    time_start = '2024-04-22 21:00:00'
    time_end = '2024-04-22 22:00:00'
    season = 'Q1'
    bridge_code = 'S263320282L0290'
    sensor_code = 'XSDDQ-DIS-G03-001-01'
    data_single = analysis_nd(bridge_code, sensor_code, time_start, time_end, bridge, season)

    # 读取 Word 文档
    doc = Document('01.塘南河桥24年4月数据分析报告.docx')
    # 指定关键词
    keyword = "混凝土结构裂缝监测"

    # 调用函数，在指定关键词的标题后添加 "hello, world"
    add_picture_to_doc_after_table(doc, keyword)



