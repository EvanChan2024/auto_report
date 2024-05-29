from docx import Document
'''
找到标题并识别标题下的第n个段落位置，添加指定话语
'''


def find_paragraph_by_keyword(doc, keyword):
    for paragraph in doc.paragraphs:
        # 查找包含指定关键词的标题段落
        if paragraph.style.name == 'Heading 2' and keyword in paragraph.text:
            return paragraph


def add_hello_world_after_second_paragraph(doc, keyword):
    # 查找包含指定关键词的标题段落
    target_paragraph = find_paragraph_by_keyword(doc, keyword)
    if target_paragraph:
        count = 0
        # 计数找到的段落数
        paragraph_count = 0
        for paragraph in doc.paragraphs:
            # 如果是目标段落之后的段落，计数加一
            if paragraph_count > 0:
                count += 1
            # 如果计数等于3，则在该段落后添加 "hello, world"
            if count == 3:
                paragraph.insert_paragraph_before("hello, world")
                break
            # 如果找到目标段落，段落计数加一
            if paragraph.text == target_paragraph.text:
                paragraph_count += 1


# 读取 Word 文档
doc = Document('03.塘南河桥数据分析报告0719.docx')

# 指定关键词
keyword = "主梁竖向位移"

# 调用函数，在指定关键词的标题后添加 "hello, world"
add_hello_world_after_second_paragraph(doc, keyword)

# 保存修改后的文档
doc.save('modified_document.docx')
